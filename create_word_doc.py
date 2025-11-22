from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create a new Document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_paragraph()
title_run = title.add_run('Spring Boot E-Commerce Backend')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('Complete Architecture Guide')
subtitle_run.font.size = Pt(18)
subtitle_run.font.italic = True
subtitle_run.font.color.rgb = RGBColor(51, 102, 153)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Space

# Table of Contents
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Overview',
    '2. Architecture & Flow',
    '3. Technology Stack',
    '4. How Components Connect',
    '5. Key Annotations Explained',
    '6. Database & ORM Concepts',
    '7. API Endpoints',
    '8. Request-Response Lifecycle'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ===== SECTION 1: OVERVIEW =====
doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'Your project is a REST API backend for an e-commerce application that allows users to:'
)
overview_items = [
    'View all products',
    'View a single product',
    'Upload products with images',
    'Retrieve product images'
]
for item in overview_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('The architecture follows the Layered/Tier Architecture pattern:')

# Architecture diagram
arch_text = doc.add_paragraph()
arch_text.add_run('Client (Browser/Mobile App)\n').font.bold = True
doc.add_paragraph('↓', style='List Bullet')
arch_text = doc.add_paragraph()
arch_text.add_run('Controller Layer').font.bold = True
arch_text.add_run(' ← Handles HTTP requests/responses')
doc.add_paragraph('↓', style='List Bullet')
arch_text = doc.add_paragraph()
arch_text.add_run('Service Layer').font.bold = True
arch_text.add_run(' ← Business logic')
doc.add_paragraph('↓', style='List Bullet')
arch_text = doc.add_paragraph()
arch_text.add_run('Repository Layer').font.bold = True
arch_text.add_run(' ← Database access')
doc.add_paragraph('↓', style='List Bullet')
arch_text = doc.add_paragraph()
arch_text.add_run('Database (H2)').font.bold = True
arch_text.add_run(' ← Data storage')

doc.add_page_break()

# ===== SECTION 2: ARCHITECTURE & FLOW =====
doc.add_heading('2. Architecture & Flow', level=1)

# Layer 1
doc.add_heading('Layer 1: Controller Layer (ProductController.java)', level=2)
doc.add_heading('Purpose:', level=3)
doc.add_paragraph('Handle incoming HTTP requests and send responses back to clients.')

doc.add_heading('How it works:', level=3)
controller_points = [
    'Receives HTTP requests (GET, POST, etc.)',
    'Parses request data (path variables, request body, file uploads)',
    'Calls business logic from the Service layer',
    'Sends HTTP responses back (JSON, images, etc.)',
    'Handles HTTP status codes (200 OK, 201 CREATED, 404 NOT FOUND, 500 ERROR)'
]
for point in controller_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('Your Controllers:', level=3)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'HTTP Method & Path'
table.rows[0].cells[1].text = 'Purpose'
table.rows[1].cells[0].text = 'GET /api/products'
table.rows[1].cells[1].text = 'Get all products'
table.rows[2].cells[0].text = 'GET /api/product/{id}'
table.rows[2].cells[1].text = 'Get single product by ID'
table.rows[3].cells[0].text = 'POST /api/product'
table.rows[3].cells[1].text = 'Add new product with image'
table.rows[4].cells[0].text = 'GET /api/product/{productId}/image'
table.rows[4].cells[1].text = 'Download product image'

doc.add_paragraph()

# Layer 2
doc.add_heading('Layer 2: Service Layer (ProductService.java)', level=2)
doc.add_heading('Purpose:', level=3)
doc.add_paragraph('Contains business logic, data validation, and orchestration.')

doc.add_heading('How it works:', level=3)
service_points = [
    'Receives data from the Controller',
    'Processes the business logic (e.g., converting image file to bytes)',
    'Validates data before saving',
    'Calls the Repository to access the database',
    'Returns processed data back to the Controller'
]
for point in service_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('Your Service Methods:', level=3)
service_methods = [
    'getAllProducts() → Retrieves all products from database',
    'getProductById(int id) → Retrieves a specific product',
    'addProduct(product, file) → Saves product + converts image to binary'
]
for method in service_methods:
    doc.add_paragraph(method, style='List Bullet')

# Layer 3
doc.add_heading('Layer 3: Repository Layer (ProductRepo.java)', level=2)
doc.add_heading('Purpose:', level=3)
doc.add_paragraph('Database access abstraction. Uses Spring Data JPA.')

doc.add_heading('How it works:', level=3)
repo_points = [
    'Extends JpaRepository<Product, Integer>',
    'Provides pre-built methods: findAll(), findById(), save(), delete(), etc.',
    'No need to write SQL queries for basic CRUD operations',
    'Spring automatically generates database queries'
]
for point in repo_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('What JpaRepository provides (out of the box):', level=3)
jpa_methods = [
    'save(entity) → INSERT or UPDATE',
    'findAll() → SELECT *',
    'findById(id) → SELECT WHERE id = ?',
    'delete(entity) → DELETE',
    'count() → COUNT(*)',
    '... and many more'
]
for method in jpa_methods:
    doc.add_paragraph(method, style='List Bullet')

# Layer 4 & 5
doc.add_heading('Layer 4: Model/Entity Layer (Product.java)', level=2)
doc.add_heading('Purpose:', level=3)
doc.add_paragraph('Defines the database table structure as a Java class.')

doc.add_heading('How it works:', level=3)
model_points = [
    'Each class with @Entity annotation = one database table',
    'Each field with @Id = primary key',
    'Each private field = one database column',
    'JPA/Hibernate automatically creates/updates the table'
]
for point in model_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('Layer 5: Database (H2 In-Memory)', level=2)
doc.add_heading('Purpose:', level=3)
doc.add_paragraph('Stores all data persistently (during runtime).')

doc.add_heading('Current Setup:', level=3)
db_setup = [
    'H2 embedded database (in-memory)',
    'URL: jdbc:h2:mem:ecomdb',
    'Automatically creates tables when application starts',
    'Data is lost when application stops'
]
for item in db_setup:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ===== SECTION 3: TECHNOLOGY STACK =====
doc.add_heading('3. Technology Stack', level=1)
doc.add_heading('Dependencies (from pom.xml):', level=2)

dependencies = [
    {
        'name': 'spring-boot-starter-data-jpa',
        'desc': 'Provides JPA (Java Persistence API) interface. Enables Hibernate ORM (Object-Relational Mapping). Converts Java objects to database records.'
    },
    {
        'name': 'spring-boot-starter-web',
        'desc': 'Provides REST API capabilities. Includes embedded Tomcat server. Handles HTTP requests/responses. Makes your application a web server.'
    },
    {
        'name': 'spring-boot-devtools',
        'desc': 'Auto-restarts application when files change. Faster development cycle.'
    },
    {
        'name': 'h2 (H2 Database)',
        'desc': 'Lightweight, in-memory database. Good for learning and testing. Stores your product data.'
    },
    {
        'name': 'lombok',
        'desc': 'Reduces boilerplate Java code. Auto-generates getters, setters, constructors. Annotations: @Data, @AllArgsConstructor, @NoArgsConstructor'
    }
]

for dep in dependencies:
    p = doc.add_paragraph()
    p_run = p.add_run(dep['name'])
    p_run.bold = True
    doc.add_paragraph(dep['desc'], style='List Bullet')

doc.add_page_break()

# ===== SECTION 4: HOW COMPONENTS CONNECT =====
doc.add_heading('4. How Components Connect', level=1)

doc.add_heading('Example 1: GET /api/products Request Flow', level=2)

flow_steps = [
    ('CLIENT sends', 'GET http://localhost:8080/api/products'),
    ('TOMCAT SERVLET', 'embedded web server receives request'),
    ('SPRING DISPATCHER', 'detects @GetMapping("/products")'),
    ('CONTROLLER', 'executes: getAllProducts()'),
    ('SERVICE', 'executes: getAllProducts() → calls repo.findAll()'),
    ('REPOSITORY', 'generates SQL: SELECT * FROM product'),
    ('DATABASE (H2)', 'returns: List of all Product records'),
    ('REPOSITORY', 'converts SQL results → List<Product> objects'),
    ('SERVICE', 'returns this List to Controller'),
    ('CONTROLLER', 'wraps in ResponseEntity with HTTP 200 OK'),
    ('SPRING JACKSON', 'converts List<Product> → JSON'),
    ('CLIENT', 'receives JSON response with all products')
]

for step_num, (component, action) in enumerate(flow_steps, 1):
    p = doc.add_paragraph()
    p_run = p.add_run(f'{step_num}. {component}: ')
    p_run.bold = True
    p.add_run(action)

doc.add_paragraph()
doc.add_heading('Example 2: POST /api/product Request Flow (with Image Upload)', level=2)

flow_steps_2 = [
    ('CLIENT sends', 'POST http://localhost:8080/api/product with FormData (product + imageFile)'),
    ('SPRING MULTIPART HANDLER', 'parses the FormData'),
    ('CONTROLLER', 'executes: addProduct(Product, MultipartFile)'),
    ('SERVICE', 'extracts image metadata (filename, content-type, bytes)'),
    ('SERVICE', 'sets image data on Product object'),
    ('SERVICE', 'calls repo.save(product)'),
    ('REPOSITORY', 'generates SQL: INSERT INTO product (...)'),
    ('DATABASE', 'stores the Product record and auto-generates ID'),
    ('REPOSITORY', 'returns saved Product object with ID'),
    ('SERVICE', 'returns Product to Controller'),
    ('CONTROLLER', 'returns ResponseEntity with HTTP 201 CREATED'),
    ('JACKSON', 'converts Product to JSON (imageData excluded by @JsonIgnore)'),
    ('CLIENT', 'receives HTTP 201 with created product data')
]

for step_num, (component, action) in enumerate(flow_steps_2, 1):
    p = doc.add_paragraph()
    p_run = p.add_run(f'{step_num}. {component}: ')
    p_run.bold = True
    p.add_run(action)

doc.add_page_break()

# ===== SECTION 5: KEY ANNOTATIONS =====
doc.add_heading('5. Key Annotations Explained', level=1)

doc.add_heading('Controller Annotations', level=2)

annotations = [
    {
        'name': '@RestController',
        'code': '@RestController\npublic class ProductController { ... }',
        'explanation': 'Tells Spring: "This class handles REST API requests". Every method returns JSON/data automatically (not HTML templates). Combines @Controller + @ResponseBody'
    },
    {
        'name': '@RequestMapping("/api")',
        'code': '@RequestMapping("/api")\npublic class ProductController { ... }',
        'explanation': 'Base URL prefix for all endpoints in this controller. All endpoints start with /api/. Example: Full URL becomes /api/products'
    },
    {
        'name': '@CrossOrigin',
        'code': '@CrossOrigin\npublic class ProductController { ... }',
        'explanation': 'Allows requests from different domains/ports. Without this: Frontend on port 3000 cannot call backend on port 8080. CORS = Cross-Origin Resource Sharing'
    },
    {
        'name': '@GetMapping("/products")',
        'code': '@GetMapping("/products")\npublic ResponseEntity<List<Product>> getAllProducts() { ... }',
        'explanation': 'Maps HTTP GET requests to /api/products to this method. Similar: @PostMapping, @PutMapping, @DeleteMapping'
    },
    {
        'name': '@PostMapping(value = "/product", consumes = {"multipart/form-data"})',
        'code': '@PostMapping(value = "/product", consumes = {"multipart/form-data"})\npublic ResponseEntity<?> addProduct(...) { ... }',
        'explanation': 'Maps HTTP POST requests to /api/product. consumes = {"multipart/form-data"} means: Accept file uploads (FormData). Without this: Cannot accept image files'
    },
    {
        'name': '@PathVariable',
        'code': '@GetMapping("/product/{id}")\npublic ResponseEntity<Product> getProduct(@PathVariable int id) { ... }',
        'explanation': 'Extracts URL parameter {id} as a Java variable. Example: /product/5 → id = 5'
    },
    {
        'name': '@RequestPart',
        'code': '@PostMapping("/product")\npublic ResponseEntity<?> addProduct(\n    @RequestPart Product product,\n    @RequestPart MultipartFile imageFile\n) { ... }',
        'explanation': 'Extracts named form fields from multipart request. @RequestPart Product → Extract form field named "product". @RequestPart MultipartFile → Extract file upload named "imageFile"'
    }
]

for ann in annotations:
    doc.add_heading(ann['name'], level=3)
    p = doc.add_paragraph()
    p_run = p.add_run('Code:')
    p_run.bold = True
    code_para = doc.add_paragraph(ann['code'], style='List Bullet')
    code_para.style.font.name = 'Courier New'

    p = doc.add_paragraph()
    p_run = p.add_run('Explanation:')
    p_run.bold = True
    doc.add_paragraph(ann['explanation'], style='List Bullet')

doc.add_heading('Entity/Model Annotations', level=2)

entity_annotations = [
    {
        'name': '@Entity',
        'explanation': 'Tells JPA: "This class represents a database table". Table name = class name (lowercase by default). Product class → product table'
    },
    {
        'name': '@Id',
        'explanation': 'Marks this field as PRIMARY KEY. Must be unique for each record. Used to uniquely identify a product.'
    },
    {
        'name': '@GeneratedValue(strategy = GenerationType.IDENTITY)',
        'explanation': 'Auto-generates ID value when new product is inserted. IDENTITY = Database auto-increment (incrementing numbers). You don\'t set ID manually; database does it.'
    },
    {
        'name': '@Lob (Large Object)',
        'explanation': 'Tells JPA: This field stores large binary data. byte[] = array of bytes (binary image data). Stores in BLOB column in database.'
    },
    {
        'name': '@Basic(fetch = FetchType.LAZY)',
        'explanation': 'LAZY = Don\'t load this field by default. When you fetch a Product, imageData is NOT loaded (saves memory/bandwidth). Load imageData only when explicitly accessed. Good for large binary data.'
    },
    {
        'name': '@JsonIgnore',
        'explanation': 'When converting Product to JSON: skip this field. Why: Don\'t send huge binary data in API response. Instead: Provide separate /image endpoint to download.'
    },
    {
        'name': '@JsonIgnoreProperties',
        'explanation': 'Ignores Hibernate proxy fields when serializing to JSON. Prevents JSON serialization errors.'
    },
    {
        'name': '@Data (Lombok)',
        'explanation': 'Auto-generates: getters, setters, toString(), equals(), hashCode(). Reduces boilerplate code. Combines: @Getter + @Setter + @ToString + @EqualsAndHashCode'
    },
    {
        'name': '@AllArgsConstructor (Lombok)',
        'explanation': 'Auto-generates constructor with all fields as parameters. Example: new Product(id, name, brand, price, ...)'
    },
    {
        'name': '@NoArgsConstructor (Lombok)',
        'explanation': 'Auto-generates empty constructor. Example: new Product()'
    }
]

for ann in entity_annotations:
    doc.add_heading(ann['name'], level=3)
    doc.add_paragraph(ann['explanation'])

doc.add_heading('Service Annotations', level=2)

service_annotations = [
    {
        'name': '@Service',
        'explanation': 'Tells Spring: "This is a service class (business logic)". Spring automatically creates an instance (bean) of this class. Can be injected into other classes.'
    },
    {
        'name': '@Autowired',
        'explanation': 'Tells Spring: "Automatically inject an instance of ProductService". Dependency Injection: Spring finds and provides the object. Don\'t use new ProductService() manually.'
    }
]

for ann in service_annotations:
    doc.add_heading(ann['name'], level=3)
    doc.add_paragraph(ann['explanation'])

doc.add_heading('Repository Annotations', level=2)
doc.add_heading('@Repository', level=3)
doc.add_paragraph('Tells Spring: "This is a data access object". Enables automatic SQL generation. <Product, Integer> = Entity type, Primary Key type.')

doc.add_page_break()

# ===== SECTION 6: DATABASE & ORM =====
doc.add_heading('6. Database & ORM Concepts', level=1)

doc.add_heading('What is ORM (Object-Relational Mapping)?', level=2)
doc.add_paragraph('ORM bridges Java objects and database tables:')

doc.add_paragraph()
orm_points = [
    'Product class ↔ product table',
    'Product object ↔ product row',
    'id field ↔ id column',
    'name field ↔ name column',
    'imageData field ↔ image_data column'
]
for point in orm_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('Benefits:', level=3)
benefits = [
    'Write Java code instead of SQL queries',
    'Database agnostic (switch from H2 to MySQL easily)',
    'Automatic type conversion (Java types ↔ SQL types)'
]
for benefit in benefits:
    doc.add_paragraph(benefit, style='List Bullet')

doc.add_heading('What is Hibernate?', level=2)
doc.add_paragraph('Hibernate is the ORM framework that:')
hibernate_points = [
    'Generates SQL queries from your Java code',
    'Converts database records to Java objects',
    'Manages entity lifecycle (new, managed, detached, removed)',
    'Handles relationships between entities'
]
for point in hibernate_points:
    doc.add_paragraph(point, style='List Number')

doc.add_heading('Example:', level=3)
doc.add_paragraph('Instead of writing:')
sql_example = doc.add_paragraph('SELECT * FROM product WHERE id = 1;', style='List Bullet')
sql_example.style.font.name = 'Courier New'

doc.add_paragraph('You write:')
java_example = doc.add_paragraph('Product product = repo.findById(1).orElse(null);', style='List Bullet')
java_example.style.font.name = 'Courier New'

doc.add_paragraph('Hibernate generates the SQL for you!')

doc.add_heading('JPA vs Hibernate:', level=2)
jpa_points = [
    'JPA = Standard interface/specification (like a contract)',
    'Hibernate = Implementation of JPA (the actual implementation)',
    'Spring Data JPA = Wrapper around Hibernate that makes it even easier'
]
for point in jpa_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('Database Table Structure (Auto-generated)', level=2)
doc.add_paragraph('Your Product class creates this table in H2:')

table_sql = doc.add_paragraph(
    'CREATE TABLE product (\n'
    '    id              INT PRIMARY KEY AUTO_INCREMENT,\n'
    '    name            VARCHAR(255),\n'
    '    desc            VARCHAR(255),\n'
    '    brand           VARCHAR(255),\n'
    '    price           DECIMAL(19,2),\n'
    '    category        VARCHAR(255),\n'
    '    release_date    TIMESTAMP,\n'
    '    available       BOOLEAN,\n'
    '    quantity        INT,\n'
    '    image_name      VARCHAR(255),\n'
    '    image_type      VARCHAR(255),\n'
    '    image_data      BLOB\n'
    ');',
    style='List Bullet'
)
table_sql.style.font.name = 'Courier New'

doc.add_heading('How:', level=3)
how_points = [
    'spring.jpa.hibernate.ddl-auto=update tells Hibernate to automatically create/update tables',
    'Field names map to column names (camelCase → snake_case)'
]
for point in how_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_page_break()

# ===== SECTION 7: API ENDPOINTS =====
doc.add_heading('7. API Endpoints', level=1)

# Endpoint 1
doc.add_heading('Endpoint 1: GET /api/products', level=2)
doc.add_heading('Purpose:', level=3)
doc.add_paragraph('Retrieve all products')

doc.add_heading('Request:', level=3)
req1 = doc.add_paragraph('GET /api/products HTTP/1.1\nHost: localhost:8080')
req1.style.font.name = 'Courier New'

doc.add_heading('Response (HTTP 200 OK):', level=3)
doc.add_paragraph('Returns a JSON array of all products with their details (excluding imageData)')

doc.add_heading('Backend Flow:', level=3)
flow1 = [
    'Controller receives GET request',
    'Calls service.getAllProducts()',
    'Service calls repo.findAll()',
    'Repository generates: SELECT * FROM product',
    'Returns List<Product>',
    'Jackson converts to JSON',
    'Returns HTTP 200 OK with JSON body'
]
for f in flow1:
    doc.add_paragraph(f, style='List Number')

# Endpoint 2
doc.add_heading('Endpoint 2: GET /api/product/{id}', level=2)
doc.add_heading('Purpose:', level=3)
doc.add_paragraph('Retrieve a single product by ID')

doc.add_heading('Request:', level=3)
req2 = doc.add_paragraph('GET /api/product/1 HTTP/1.1\nHost: localhost:8080')
req2.style.font.name = 'Courier New'

doc.add_heading('Response:', level=3)
doc.add_paragraph('HTTP 200 OK: Returns the product as JSON')
doc.add_paragraph('HTTP 404 Not Found: If product does not exist')

doc.add_heading('Backend Flow:', level=3)
flow2 = [
    '@PathVariable int id extracts 1 from URL',
    'Controller calls service.getProductById(1)',
    'Service calls repo.findById(1)',
    'Repository generates: SELECT * FROM product WHERE id = 1',
    'If found: returns Product object → HTTP 200 OK',
    'If NOT found: returns null → HTTP 404 Not Found'
]
for f in flow2:
    doc.add_paragraph(f, style='List Number')

# Endpoint 3
doc.add_heading('Endpoint 3: POST /api/product', level=2)
doc.add_heading('Purpose:', level=3)
doc.add_paragraph('Add new product with image')

doc.add_heading('Request:', level=3)
req3 = doc.add_paragraph(
    'POST /api/product HTTP/1.1\n'
    'Content-Type: multipart/form-data\n\n'
    'FormData:\n'
    '- product: { "name": "Laptop", "brand": "Dell", "price": 79999.99, ... }\n'
    '- imageFile: <binary image data>'
)
req3.style.font.name = 'Courier New'

doc.add_heading('Response (HTTP 201 CREATED):', level=3)
doc.add_paragraph('Returns created product with auto-generated ID')

doc.add_heading('Backend Flow:', level=3)
flow3 = [
    '@RequestPart Product product deserializes JSON → Product object',
    '@RequestPart MultipartFile imageFile receives uploaded file',
    'Controller calls service.addProduct(product, imageFile)',
    'Service extracts image metadata and converts to bytes',
    'Service calls repo.save(product)',
    'Repository generates INSERT query',
    'Database assigns auto-generated ID',
    'Returns saved Product with ID',
    'HTTP 201 CREATED response'
]
for f in flow3:
    doc.add_paragraph(f, style='List Number')

# Endpoint 4
doc.add_heading('Endpoint 4: GET /api/product/{productId}/image', level=2)
doc.add_heading('Purpose:', level=3)
doc.add_paragraph('Download product image')

doc.add_heading('Request:', level=3)
req4 = doc.add_paragraph('GET /api/product/1/image HTTP/1.1\nHost: localhost:8080')
req4.style.font.name = 'Courier New'

doc.add_heading('Response (HTTP 200 OK):', level=3)
doc.add_paragraph('Binary image data with appropriate Content-Type header')

doc.add_heading('Backend Flow:', level=3)
flow4 = [
    '@PathVariable int productId extracts 1 from URL',
    'Controller calls service.getProductById(1)',
    'Service queries database',
    'Controller extracts: product.getImageData() (byte array)',
    'Controller sets HTTP response header: Content-Type: image/jpeg',
    'Returns binary data',
    'Browser displays/downloads image'
]
for f in flow4:
    doc.add_paragraph(f, style='List Number')

doc.add_page_break()

# ===== SECTION 8: REQUEST-RESPONSE LIFECYCLE =====
doc.add_heading('8. Request-Response Lifecycle', level=1)

doc.add_heading('Complete Lifecycle Example: Create Product with Image', level=2)

lifecycle_steps = [
    {
        'step': 'CLIENT Sends Request',
        'details': 'User sends POST /api/product with FormData (product JSON + image file)'
    },
    {
        'step': 'Spring Web Server (Tomcat) Receives Request',
        'details': 'Tomcat identifies HTTP Method (POST), URL (/api/product), and body (Multipart form data)'
    },
    {
        'step': 'Spring Dispatcher Maps to Controller',
        'details': 'Spring finds @PostMapping("/product") in ProductController'
    },
    {
        'step': 'Multipart Parser Extracts Data',
        'details': 'Extracts "product" field (JSON deserized to Product object) and "imageFile" field (MultipartFile wrapper)'
    },
    {
        'step': 'Controller Method Executes',
        'details': 'ProductController.addProduct(Product, MultipartFile) calls service.addProduct(product, imageFile)'
    },
    {
        'step': 'Service Layer - Business Logic',
        'details': 'Extracts filename, content-type, and bytes from file. Sets these on Product. Calls repo.save(product)'
    },
    {
        'step': 'Repository Layer - Database Access',
        'details': 'Hibernate generates: INSERT INTO product (name, brand, price, image_name, image_type, image_data, ...) VALUES (...)'
    },
    {
        'step': 'Database Execution',
        'details': 'H2 executes INSERT, auto-generates ID, stores row, returns success'
    },
    {
        'step': 'Return Data to Service',
        'details': 'Hibernate converts database row to Product Java object with auto-generated ID'
    },
    {
        'step': 'Service Returns to Controller',
        'details': 'Service returns Product { id: 1, name: "Laptop", brand: "Dell", ... }'
    },
    {
        'step': 'Controller Wraps Response',
        'details': 'Controller returns ResponseEntity.status(201).body(savedProduct). HTTP Status: 201 CREATED'
    },
    {
        'step': 'Jackson Serializes to JSON',
        'details': 'Jackson converts Product object to JSON (imageData excluded by @JsonIgnore)'
    },
    {
        'step': 'HTTP Response Sent to Client',
        'details': 'HTTP/1.1 201 Created with JSON body'
    },
    {
        'step': 'Client Receives Response',
        'details': 'Frontend receives HTTP Status 201 with created product data'
    }
]

for num, ls in enumerate(lifecycle_steps, 1):
    p = doc.add_paragraph()
    p_run = p.add_run(f'Step {num}: {ls["step"]}')
    p_run.bold = True
    p_run.font.size = Pt(12)
    doc.add_paragraph(ls['details'], style='List Bullet')

doc.add_page_break()

# ===== SUMMARY =====
doc.add_heading('Summary: How Everything Works Together', level=1)

doc.add_heading('Request Path (What happens when you call an API):', level=2)

request_path = [
    'HTTP Request',
    'Tomcat (Web Server)',
    'Spring Dispatcher (Router)',
    'Controller (Request Handler)',
    'Service (Business Logic)',
    'Repository (Database Access)',
    'Hibernate (ORM - Object to SQL)',
    'H2 Database (SQL Execution & Storage)',
    '[Data returned back up the chain]',
    'Jackson (Object to JSON)',
    'HTTP Response',
    'Client (Browser/App)'
]

for path in request_path:
    if path.startswith('['):
        doc.add_paragraph(path, style='List Bullet')
    else:
        doc.add_paragraph(path, style='List Bullet')

doc.add_heading('Key Concepts Recap:', level=2)

concepts = [
    ('@RestController', 'Handles REST requests'),
    ('@RequestMapping', 'URL prefix'),
    ('@GetMapping/@PostMapping', 'HTTP method + URL routing'),
    ('@Autowired', 'Dependency Injection (automatic wiring)'),
    ('@Entity', 'Database table'),
    ('@Id @GeneratedValue', 'Primary Key with auto-increment'),
    ('@Lob', 'Large binary data (images)'),
    ('@Lazy', 'Load data only when needed'),
    ('@JsonIgnore', 'Don\'t include in JSON response'),
    ('JpaRepository', 'Pre-built database access methods'),
    ('Service', 'Business logic layer'),
    ('DTO', 'Data Transfer Object (if needed for specific responses)')
]

for concept, explanation in concepts:
    p = doc.add_paragraph()
    p_run = p.add_run(concept)
    p_run.bold = True
    p.add_run(f' = {explanation}')

doc.add_page_break()

# ===== NEXT STEPS =====
doc.add_heading('Next Steps for Learning', level=1)
doc.add_paragraph('Now that you understand the architecture, you can enhance your project by implementing:')

next_steps = [
    'Add a DELETE endpoint (@DeleteMapping) to remove products',
    'Add an UPDATE endpoint (@PutMapping) to modify products',
    'Add custom search methods in ProductRepo to find products by category/brand',
    'Add category filter to getAllProducts for advanced filtering',
    'Implement pagination using Page<Product> instead of List<Product>',
    'Add data validation using @Valid and validation annotations',
    'Create custom exceptions and error handlers for better error management',
    'Add unit tests for service and controller methods',
    'Implement role-based access control (authentication & authorization)',
    'Add API documentation using Swagger/SpringFox'
]

for num, step in enumerate(next_steps, 1):
    doc.add_paragraph(step, style='List Number')

# Final note
final_note = doc.add_paragraph()
final_note_run = final_note.add_run('Your project is now ready for these enhancements!')
final_note_run.bold = True
final_note_run.font.size = Pt(12)
final_note_run.font.color.rgb = RGBColor(0, 102, 0)

# Save the document
output_path = r'C:\Users\z00542kh\Desktop\ecom-proj\Spring_Boot_Architecture_Guide.docx'
doc.save(output_path)

print(f"Word document created successfully!")
print(f"Location: {output_path}")

